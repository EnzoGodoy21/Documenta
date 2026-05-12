"""
gera_word.py — Gerador de documentos Word por tabela

Placeholders suportados no template:
  [TXT:NOME_TABELA]  → substituído pelo nome da tabela (texto)
  [TXT:outra_chave]  → substituído pelo valor correspondente (texto genérico)
  [IMG:chave]        → substituído por prints/chave_TABELA.png (imagem)

Uso:
  python gera_word.py                            # execução normal
  python gera_word.py --force                    # reprocessa tudo
  python gera_word.py --tabelas VENDAS CLIENTES  # só essas tabelas
  python gera_word.py --workers 8                # paralelismo
  python gera_word.py --template outro.docx      # template específico
"""

import os
import re
import argparse
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def iter_paragrafos(doc):
    """Itera sobre todos os parágrafos do documento, incluindo dentro de tabelas."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _substituir_texto_no_paragrafo(p, busca, substituto):
    """
    Substitui texto em um parágrafo, tratando casos onde o placeholder
    está dividido entre múltiplos runs.
    """
    if busca not in p.text:
        return

    # Caso simples: placeholder dentro de um único run
    for run in p.runs:
        if busca in run.text:
            run.text = run.text.replace(busca, substituto)
            return

    # Caso complexo: placeholder dividido entre runs — mescla tudo no primeiro
    texto_completo = "".join(r.text for r in p.runs)
    if busca in texto_completo:
        novo_texto = texto_completo.replace(busca, substituto)
        if p.runs:
            p.runs[0].text = novo_texto
            for run in p.runs[1:]:
                run.text = ""


def _substituir_imagem_no_paragrafo(p, placeholder, caminho_imagem):
    """
    Substitui um placeholder de imagem pelo arquivo de imagem correspondente.
    Centraliza e remove indentação.
    """
    # Limpa todos os runs
    for run in p.runs:
        run.text = ""

    # Zera indentação
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    pPr.append(ind)

    # Centraliza
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insere imagem
    run = p.add_run()
    run.add_picture(caminho_imagem, width=Inches(5.5))


# ─────────────────────────────────────────────────────────────
# AUTO-DESCOBERTA
# ─────────────────────────────────────────────────────────────

def descobrir_chaves(template_path: str) -> dict:
    """
    Escaneia o template e retorna as chaves encontradas separadas por tipo.

    Retorna:
      {
        "img": ["chave1", "chave2", ...],   # [IMG:chave]
        "txt": ["chave1", "chave2", ...],   # [TXT:chave]
      }

    Exemplos:
      [IMG:visao_geral]   → img: 'visao_geral'
      [TXT:NOME_TABELA]   → txt: 'NOME_TABELA'
    """
    doc = Document(template_path)
    chaves_img = set()
    chaves_txt = set()

    for p in iter_paragrafos(doc):
        chaves_img.update(re.findall(r"\[IMG:([^\]]+)\]", p.text))
        chaves_txt.update(re.findall(r"\[TXT:([^\]]+)\]", p.text))

    return {
        "img": sorted(chaves_img),
        "txt": sorted(chaves_txt),
    }


def descobrir_tabelas(prints_path: str, chaves_img: list) -> list:
    """
    Descobre os nomes de tabelas cruzando os arquivos em prints/
    com as chaves de imagem encontradas no template.

    Convenções suportadas:
      chave_TABELA.png          (único print por chave)
      chave_1_TABELA.png        (múltiplos prints numerados por chave)

    Tabelas são inferidas apenas via [IMG:*] — não via [TXT:*].
    """
    pasta = Path(prints_path)
    if not pasta.exists():
        return []

    tabelas = set()
    extensoes = {".png", ".jpg", ".jpeg"}

    for arquivo in pasta.iterdir():
        if arquivo.suffix.lower() not in extensoes:
            continue
        nome = arquivo.stem
        for chave in chaves_img:
            prefixo = chave + "_"
            if not nome.startswith(prefixo):
                continue
            resto = nome[len(prefixo):]
            if not resto:
                continue
            # Verifica convenção numerada: chave_N_TABELA (N inteiro)
            partes = resto.split("_", 1)
            if len(partes) == 2 and partes[0].isdigit():
                tabela = partes[1]
            else:
                tabela = resto
            if tabela:
                tabelas.add(tabela)

    return sorted(tabelas)


def ler_tabelas_de_arquivo(path: str) -> tuple:
    """
    Lê nomes de tabelas de um arquivo .txt ou .csv.

    .txt — uma tabela por linha; linhas em branco e comentários (#) ignorados.
           Retorna dados vazio (sem colunas extras).
    .csv — lê a coluna 'nome_tabela' se existir; caso contrário usa a primeira coluna.
           Colunas extras (além de nome_tabela) são retornadas como dados por tabela,
           servindo como valores [TXT:*] no template.
           Suporta delimitadores vírgula (,) e ponto-e-vírgula (;).

    Retorna:
      (tabelas, dados)
        tabelas — list[str]: nomes das tabelas na ordem do arquivo
        dados   — dict[str, dict]: valores extras por tabela
                  ex.: {"VENDAS": {"RESPONSAVEL": "João", "DOMINIO": "Comercial"}}
    """
    import csv as _csv

    ext = Path(path).suffix.lower()
    tabelas = []
    dados = {}

    with open(path, encoding="utf-8-sig") as f:
        if ext == ".csv":
            amostra = f.read(4096)
            f.seek(0)
            try:
                dialeto = _csv.Sniffer().sniff(amostra, delimiters=",;\t|")
            except _csv.Error:
                # Detecta manualmente pelo delimitador mais frequente
                dialeto = type("_D", (_csv.excel,), {
                    "delimiter": ";" if amostra.count(";") >= amostra.count(",") else ","
                })
            reader = _csv.DictReader(f, dialect=dialeto)

            # Normaliza nomes de colunas (strip + uppercase)
            fieldnames_raw = reader.fieldnames or []
            fieldnames_norm = [c.strip().upper() for c in fieldnames_raw]
            col_map = dict(zip(fieldnames_norm, fieldnames_raw))  # NORM → raw

            col_norm = next((c for c in fieldnames_norm if c == "NOME_TABELA"), None)
            if col_norm is None and fieldnames_norm:
                col_norm = fieldnames_norm[0]
            col_raw = col_map.get(col_norm) if col_norm else None

            colunas_extras = [c for c in fieldnames_norm if c != col_norm]

            for row in reader:
                if col_raw:
                    val = row.get(col_raw, "").strip()
                    if val:
                        tabelas.append(val)
                        if colunas_extras:
                            dados[val] = {
                                c: row.get(col_map[c], "").strip()
                                for c in colunas_extras
                            }
        else:
            for linha in f:
                val = linha.strip()
                if val and not val.startswith("#"):
                    tabelas.append(val)

    return tabelas, dados


def resolver_template(template_arg: str) -> str:
    """
    Resolve o caminho do template — aceita arquivo direto ou pasta
    (retorna o primeiro .docx encontrado).
    """
    p = Path(template_arg)
    if p.is_file():
        return str(p)
    if p.is_dir():
        candidatos = sorted(p.glob("*.docx"))
        if candidatos:
            return str(candidatos[0])
    raise FileNotFoundError(f"Template não encontrado: {template_arg}")


# ─────────────────────────────────────────────────────────────
# PROCESSAMENTO
# ─────────────────────────────────────────────────────────────

def _builtin_textos(tabela: str, timestamp: datetime = None) -> dict:
    """
    Valores de texto automáticos disponíveis em todo documento, sem necessidade de CSV.

    Placeholder         Exemplo de valor
    ──────────────────────────────────────────────────
    [TXT:NOME_TABELA]   VENDAS
    [TXT:DATA]          04/04/2026
    [TXT:DATA_HORA]     04/04/2026 14:30
    [TXT:ANO]           2026
    [TXT:MES]           04
    [TXT:DIA]           04
    """
    ts = timestamp or datetime.now()
    return {
        "NOME_TABELA": tabela,
        "DATA":        ts.strftime("%d/%m/%Y"),
        "DATA_HORA":   ts.strftime("%d/%m/%Y %H:%M"),
        "ANO":         ts.strftime("%Y"),
        "MES":         ts.strftime("%m"),
        "DIA":         ts.strftime("%d"),
    }


def processar_tabela(tabela, template_path, prints_path, output_path, chaves,
                     textos=None, callback=None, timestamp=None, prefixo=""):
    """
    Processa uma tabela: abre o template, faz substituições e salva o .docx.

    Parâmetros:
      chaves     — dict {"img": [...], "txt": [...]} retornado por descobrir_chaves()
      textos     — dict opcional com valores vindos de CSV/dados externos
                   ex.: {"RESPONSAVEL": "Enzo", "DOMINIO": "Comercial"}
      timestamp  — datetime fixo para DATA/DATA_HORA (passado pelo caller para
                   garantir que todos os docs do mesmo lote tenham o mesmo valor).
                   Se None, usa datetime.now() no momento da chamada.

    Valores automáticos (sempre disponíveis, sem CSV):
      [TXT:NOME_TABELA]  [TXT:DATA]  [TXT:DATA_HORA]  [TXT:ANO]  [TXT:MES]  [TXT:DIA]

    Retorna um dict com:
      - tabela: nome da tabela
      - status: 'ok' | 'parcial' | 'erro'
      - prints_ok: chaves [IMG:*] inseridas com sucesso
      - prints_ausentes: chaves [IMG:*] sem arquivo correspondente
      - textos_ausentes: chaves [TXT:*] sem valor definido
      - erro: mensagem de erro (None se sem erro)
    """
    textos = textos or {}

    # Builtins primeiro, textos externos sobrescrevem se houver conflito
    valores_txt = {**_builtin_textos(tabela, timestamp), **textos}

    resultado = {
        "tabela": tabela,
        "status": "ok",
        "prints_ok": [],
        "prints_ausentes": [],
        "textos_ausentes": [],
        "erro": None,
    }

    try:
        doc = Document(template_path)

        # Converte em lista para que parágrafos inseridos dinamicamente
        # (múltiplas imagens) não sejam reprocessados no mesmo loop
        for p in list(iter_paragrafos(doc)):

            # ── substituição de texto [TXT:chave] ──────────────
            for chave in chaves.get("txt", []):
                placeholder = f"[TXT:{chave}]"
                if placeholder not in p.text:
                    continue
                if chave in valores_txt:
                    _substituir_texto_no_paragrafo(p, placeholder, valores_txt[chave])
                else:
                    # Mantém o placeholder e registra ausência
                    resultado["textos_ausentes"].append(chave)
                    resultado["status"] = "parcial"

            # ── substituição de imagens [IMG:chave] ────────────
            for chave in chaves.get("img", []):
                placeholder = f"[IMG:{chave}]"
                if placeholder not in p.text:
                    continue

                caminhos = _achar_prints(prints_path, chave, tabela)
                if caminhos:
                    # Primeiro print substitui o placeholder
                    _substituir_imagem_no_paragrafo(p, placeholder, caminhos[0])
                    resultado["prints_ok"].append(chave)
                    # Prints adicionais inseridos como novos parágrafos após o atual
                    p_atual = p
                    for caminho_extra in caminhos[1:]:
                        p_atual = _inserir_imagem_apos(p_atual, doc, caminho_extra)
                else:
                    # Mantém o placeholder visível para facilitar identificação
                    resultado["prints_ausentes"].append(chave)
                    resultado["status"] = "parcial"

        # Salva o documento
        os.makedirs(output_path, exist_ok=True)
        nome_arquivo = f"{prefixo}{tabela}.docx" if prefixo else f"{tabela}.docx"
        caminho_saida = Path(output_path) / nome_arquivo
        doc.save(str(caminho_saida))

    except Exception as e:
        resultado["status"] = "erro"
        resultado["erro"] = str(e)

    if callback:
        callback(resultado)

    return resultado


def _achar_prints(prints_path, chave, tabela):
    """
    Busca arquivos de print para chave+tabela.

    Suporta:
      chave_TABELA.png                       (único)
      chave_1_TABELA.png, chave_2_TABELA.png, ...  (múltiplos numerados)

    Retorna lista ordenada de caminhos (vazia se nenhum encontrado).
    """
    extensoes = (".png", ".jpg", ".jpeg")
    pasta = Path(prints_path)

    # Tenta arquivo único
    for ext in extensoes:
        candidato = pasta / f"{chave}_{tabela}{ext}"
        if candidato.exists():
            return [str(candidato)]

    # Tenta numerados: chave_1_TABELA, chave_2_TABELA, ...
    encontrados = []
    n = 1
    while True:
        achou = False
        for ext in extensoes:
            candidato = pasta / f"{chave}_{n}_{tabela}{ext}"
            if candidato.exists():
                encontrados.append(str(candidato))
                achou = True
                break
        if not achou:
            break
        n += 1

    return encontrados


def _inserir_imagem_apos(p_ref, doc, caminho_imagem):
    """Insere novo parágrafo com imagem centralizada logo após p_ref."""
    # Cria parágrafo temporário no final e move depois
    novo_p = doc.add_paragraph()

    pPr = novo_p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    pPr.append(ind)
    novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = novo_p.add_run()
    run.add_picture(caminho_imagem, width=Inches(5.5))

    # Move para logo após p_ref (lxml remove do local original ao re-inserir)
    p_ref._p.addnext(novo_p._p)

    return novo_p


# ─────────────────────────────────────────────────────────────
# RELATÓRIO
# ─────────────────────────────────────────────────────────────

def gerar_relatorio(resultados, logs_path) -> str:
    """Gera um relatório .txt com o resumo da execução. Retorna o caminho do arquivo."""
    os.makedirs(logs_path, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    caminho = Path(logs_path) / f"relatorio_{timestamp}.txt"

    ok      = [r for r in resultados if r["status"] == "ok"]
    parcial = [r for r in resultados if r["status"] == "parcial"]
    erro    = [r for r in resultados if r["status"] == "erro"]

    with open(caminho, "w", encoding="utf-8") as f:
        f.write(f"Relatório de execução — {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n")
        f.write("=" * 60 + "\n\n")
        f.write(f"Total processado : {len(resultados)}\n")
        f.write(f"  OK             : {len(ok)}\n")
        f.write(f"  Parcial        : {len(parcial)}\n")
        f.write(f"  Erro           : {len(erro)}\n\n")

        if parcial:
            f.write("── PARCIAIS ──\n")
            for r in parcial:
                ausencias = []
                if r["prints_ausentes"]:
                    ausencias.append(f"IMG ausentes: {', '.join(r['prints_ausentes'])}")
                if r["textos_ausentes"]:
                    ausencias.append(f"TXT ausentes: {', '.join(r['textos_ausentes'])}")
                f.write(f"  {r['tabela']}: {' | '.join(ausencias)}\n")
            f.write("\n")

        if erro:
            f.write("── ERROS ──\n")
            for r in erro:
                f.write(f"  {r['tabela']}: {r['erro']}\n")
            f.write("\n")

        f.write("── DETALHES ──\n")
        for r in sorted(resultados, key=lambda x: x["tabela"]):
            f.write(f"  [{r['status'].upper():7}] {r['tabela']}\n")

    return str(caminho)


# ─────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────

EPILOG = """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 PLACEHOLDERS DISPONÍVEIS NO TEMPLATE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  [TXT:*] — texto
  ──────────────────────────────────────────────────────────
  Automáticos (sem configuração):
    [TXT:NOME_TABELA]   nome da tabela              ex.: VENDAS
    [TXT:DATA]          data de geração             ex.: 04/04/2026
    [TXT:DATA_HORA]     data e hora de geração      ex.: 04/04/2026 14:30
    [TXT:ANO]           ano                         ex.: 2026
    [TXT:MES]           mês                         ex.: 04
    [TXT:DIA]           dia                         ex.: 04

  Via CSV (colunas extras do arquivo de tabelas):
    [TXT:RESPONSAVEL]   coluna RESPONSAVEL do CSV
    [TXT:DOMINIO]       coluna DOMINIO do CSV
    [TXT:DESCRICAO]     coluna DESCRICAO do CSV
    [TXT:coluna]        qualquer outra coluna

  [IMG:*] — imagem
  ──────────────────────────────────────────────────────────
    [IMG:chave]         prints/chave_TABELA.png (ou .jpg)

  [LEG:*] — legenda de print (reservado, em desenvolvimento)
  ──────────────────────────────────────────────────────────
    [LEG:chave]         texto abaixo de [IMG:chave], estilo próprio

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 CONVENÇÃO DE NOMES DOS PRINTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  chave_NOMETABELA.png   (ou .jpg / .jpeg)

  Exemplos:
    visao_geral_VENDAS.png
    distribuicao_CLIENTES.png

  As tabelas são descobertas automaticamente cruzando os
  arquivos em prints/ com as chaves [IMG:*] do template.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 CSV DE TABELAS — FILTRO + DADOS [TXT:*]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  O CSV serve ao mesmo tempo como filtro de execução e como
  fonte de dados [TXT:*] para cada tabela. Colunas extras
  além de nome_tabela viram automaticamente placeholders.

  Exemplo (tabelas.csv):
    nome_tabela ; RESPONSAVEL  ; DOMINIO
    VENDAS      ; João Silva   ; Comercial
    CLIENTES    ; Maria Santos ; CRM

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 EXEMPLOS DE USO
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  # Execução padrão (só tabelas sem .docx em output/)
  python gera_word.py

  # Reprocessar tudo do zero
  python gera_word.py --force

  # Filtrar por arquivo (também alimenta [TXT:*] se tiver colunas extras)
  python gera_word.py --tabelas-arquivo tabelas.csv

  # Tabelas específicas por nome
  python gera_word.py --tabelas VENDAS CLIENTES PEDIDOS

  # Template e pastas customizados
  python gera_word.py --template docs/meu_template.docx --prints capturas/ --output gerados/

  # Mais paralelismo para grandes volumes
  python gera_word.py --workers 8

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 STATUS DOS DOCUMENTOS GERADOS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  ok       Todos os placeholders substituídos
  parcial  Gerado, mas algum print ou texto estava ausente
           (placeholder visível no .docx para identificação)
  erro     Falha crítica — documento não foi gerado
"""


def main():
    parser = argparse.ArgumentParser(
        description=(
            "Gera documentos .docx em massa a partir de um template Word,\n"
            "substituindo placeholders [TXT:*] e [IMG:*] por tabela.\n"
            "Use --help para ver exemplos e instruções de template."
        ),
        epilog=EPILOG,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--template", default="template/",
                        metavar="CAMINHO",
                        help="Arquivo .docx de template ou pasta que o contém  (padrão: template/)")
    parser.add_argument("--prints",   default="prints/",
                        metavar="PASTA",
                        help="Pasta com os prints chave_TABELA.png             (padrão: prints/)")
    parser.add_argument("--output",   default="output/",
                        metavar="PASTA",
                        help="Pasta de saída dos .docx gerados                 (padrão: output/)")
    parser.add_argument("--logs",     default="logs/",
                        metavar="PASTA",
                        help="Pasta para relatórios de execução                (padrão: logs/)")
    parser.add_argument("--force",    action="store_true",
                        help="Reprocessa tabelas que já têm .docx em output/")
    parser.add_argument("--tabelas",         nargs="+", metavar="TABELA",
                        help="Processa apenas as tabelas listadas (nomes separados por espaço)")
    parser.add_argument("--tabelas-arquivo", metavar="ARQUIVO",
                        help="Arquivo .txt ou .csv com nomes de tabelas, um por linha")
    parser.add_argument("--workers",  type=int, default=4,
                        metavar="N",
                        help="Número de workers paralelos                      (padrão: 4)")
    parser.add_argument("--prefixo",  default="",
                        metavar="TEXTO",
                        help="Prefixo no nome dos arquivos gerados              ex.: DOC_ → DOC_VENDAS.docx")
    args = parser.parse_args()

    # Resolve template
    try:
        template_path = resolver_template(args.template)
    except FileNotFoundError as e:
        print(f"❌ {e}")
        return

    print(f"📄 Template : {template_path}")

    # Auto-descoberta de chaves
    chaves = descobrir_chaves(template_path)

    print(f"🖼️  IMG chaves: {chaves['img'] if chaves['img'] else '(nenhuma)'}")
    print(f"📝 TXT chaves: {chaves['txt'] if chaves['txt'] else '(nenhuma)'}")

    # Fonte de tabelas
    dados_csv: dict = {}
    tabelas_explicitas: list = []

    if args.tabelas_arquivo:
        try:
            do_arquivo, dados_csv = ler_tabelas_de_arquivo(args.tabelas_arquivo)
            tabelas_explicitas.extend(do_arquivo)
            extras = len(next(iter(dados_csv.values()), {}).keys()) if dados_csv else 0
            info_extras = f", {extras} coluna(s) extra(s)" if extras else ""
            print(f"📂 Arquivo  : {len(do_arquivo)} tabela(s) carregada(s) de {args.tabelas_arquivo}{info_extras}")
        except Exception as e:
            print(f"❌ Erro ao ler arquivo de tabelas: {e}")
            return

    if args.tabelas:
        for t in args.tabelas:
            if t not in tabelas_explicitas:
                tabelas_explicitas.append(t)

    if tabelas_explicitas:
        # Lista explícita fornecida — usa diretamente (não depende de prints)
        tabelas = tabelas_explicitas
        print(f"📊 Tabelas  : {len(tabelas)} (lista explícita)")
    else:
        # Sem lista — descobre via cruzamento prints/ × [IMG:*]
        tabelas = descobrir_tabelas(args.prints, chaves["img"])
        print(f"📊 Tabelas  : {len(tabelas)} encontrada(s) via prints/")

    prefixo = args.prefixo or ""
    if prefixo:
        print(f"🏷️  Prefixo  : {prefixo}")

    if not args.force:
        pendentes  = [t for t in tabelas if not (Path(args.output) / f"{prefixo}{t}.docx").exists()]
        ignoradas  = len(tabelas) - len(pendentes)
        if ignoradas:
            print(f"⏭️  Ignoradas: {ignoradas} já processada(s) (use --force para reprocessar)")
        tabelas = pendentes

    if not tabelas:
        print("✅ Nada a processar.")
        return

    # Timestamp fixo para todos os docs do lote — DATA/DATA_HORA consistentes
    ts = datetime.now()
    print(f"\n🚀 Processando {len(tabelas)} tabela(s) com {args.workers} worker(s)...")
    print(f"📅 Timestamp : {ts.strftime('%d/%m/%Y %H:%M:%S')}\n")

    resultados = []

    def on_done(resultado):
        resultados.append(resultado)
        icon = {"ok": "✅", "parcial": "⚠️ ", "erro": "❌"}.get(resultado["status"], "?")
        linha = f"{icon} [{resultado['status'].upper():7}] {resultado['tabela']}"
        if resultado["prints_ausentes"]:
            linha += f"  — IMG ausentes: {', '.join(resultado['prints_ausentes'])}"
        if resultado["textos_ausentes"]:
            linha += f"  — TXT ausentes: {', '.join(resultado['textos_ausentes'])}"
        if resultado["erro"]:
            linha += f"  — erro: {resultado['erro']}"
        print(linha)

    with ThreadPoolExecutor(max_workers=args.workers) as executor:
        futuros = {
            executor.submit(processar_tabela, t, template_path, args.prints,
                            args.output, chaves, dados_csv.get(t), on_done, ts,
                            prefixo): t
            for t in tabelas
        }
        for f in as_completed(futuros):
            pass  # resultados tratados no callback

    caminho_rel = gerar_relatorio(resultados, args.logs)

    ok      = sum(1 for r in resultados if r["status"] == "ok")
    parcial = sum(1 for r in resultados if r["status"] == "parcial")
    erro    = sum(1 for r in resultados if r["status"] == "erro")

    print(f"\n{'=' * 40}")
    print(f"✅ OK: {ok}  ⚠️  Parcial: {parcial}  ❌ Erro: {erro}")
    print(f"📋 Relatório: {caminho_rel}")


if __name__ == "__main__":
    main()
